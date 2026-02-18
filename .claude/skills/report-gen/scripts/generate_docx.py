"""Generate a formatted .docx report from structured JSON input.

Usage:
    python generate_docx.py --input report-content.json --output report.docx
"""

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# --- Theme constants ---
DEEP_BLUE = RGBColor(0x1A, 0x3C, 0x6E)
ACCENT_BLUE = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BLUE_BG = RGBColor(0xD6, 0xE4, 0xF0)
DARK_TEXT = RGBColor(0x33, 0x33, 0x33)
GRAY_TEXT = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = RGBColor(0x2E, 0x75, 0xB6)
TABLE_ALT_ROW = RGBColor(0xF2, 0xF7, 0xFC)
FONT_NAME = "Calibri"


def set_cell_shading(cell, color: RGBColor):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): f"{color}",
    })
    shading.append(shd)


def style_document(doc: Document):
    """Apply global styles to the document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    font.color.rgb = DARK_TEXT

    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # Heading styles
    for level, (size, color) in enumerate(
        [(22, DEEP_BLUE), (16, DEEP_BLUE), (13, ACCENT_BLUE)], start=1
    ):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = FONT_NAME
        heading_style.font.size = Pt(size)
        heading_style.font.color.rgb = color
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        heading_style.paragraph_format.space_after = Pt(6)


def add_heading(doc: Document, text: str, level: int = 1):
    """Add a heading paragraph."""
    doc.add_heading(text, level=level)


def add_subtitle(doc: Document, text: str):
    """Add a subtitle (smaller, gray, italic)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY_TEXT
    run.font.italic = True
    p.paragraph_format.space_after = Pt(10)


def add_paragraph(doc: Document, text: str):
    """Add a normal paragraph."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def add_bullet_list(doc: Document, items: list[str]):
    """Add a bulleted list."""
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered_list(doc: Document, items: list[str]):
    """Add a numbered list."""
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """Add a formatted table with header row styling."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)
                run.font.name = FONT_NAME
        set_cell_shading(hdr_cells[i], TABLE_HEADER_BG)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = FONT_NAME
        # Alternate row shading
        if row_idx % 2 == 1:
            for cell in row_cells:
                set_cell_shading(cell, TABLE_ALT_ROW)

    # Add spacing after table
    doc.add_paragraph()


def add_key_value(doc: Document, pairs: list[dict]):
    """Add key-value pairs as a compact two-column table."""
    table = doc.add_table(rows=0, cols=2)
    table.autofit = True
    for pair in pairs:
        row = table.add_row()
        key_cell = row.cells[0]
        val_cell = row.cells[1]
        key_cell.text = str(pair.get("key", ""))
        val_cell.text = str(pair.get("value", ""))
        for p in key_cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT_NAME
        for p in val_cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = FONT_NAME
    doc.add_paragraph()


def add_separator(doc: Document):
    """Add a horizontal line separator."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "6",
        qn("w:space"): "1",
        qn("w:color"): f"{ACCENT_BLUE}",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


# --- Section dispatcher ---
SECTION_HANDLERS = {
    "heading1": lambda doc, s: add_heading(doc, s["text"], 1),
    "heading2": lambda doc, s: add_heading(doc, s["text"], 2),
    "heading3": lambda doc, s: add_heading(doc, s["text"], 3),
    "subtitle": lambda doc, s: add_subtitle(doc, s["text"]),
    "paragraph": lambda doc, s: add_paragraph(doc, s["text"]),
    "bullet_list": lambda doc, s: add_bullet_list(doc, s["items"]),
    "numbered_list": lambda doc, s: add_numbered_list(doc, s["items"]),
    "table": lambda doc, s: add_table(doc, s["headers"], s["rows"]),
    "key_value": lambda doc, s: add_key_value(doc, s["pairs"]),
    "separator": lambda doc, s: add_separator(doc),
}


def generate_docx(data: dict, output_path: Path):
    """Generate a .docx file from structured JSON data."""
    doc = Document()
    style_document(doc)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Process each section
    for section in data.get("sections", []):
        section_type = section.get("type", "")
        handler = SECTION_HANDLERS.get(section_type)
        if handler:
            handler(doc, section)
        else:
            print(f"Warning: unknown section type '{section_type}', skipping")

    # Footer
    footer_text = data.get("footer", "")
    if footer_text:
        add_separator(doc)
        p = doc.add_paragraph()
        run = p.add_run(footer_text)
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY_TEXT
        run.font.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Generated: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Generate .docx from JSON")
    parser.add_argument("--input", required=True, help="Input JSON file path")
    parser.add_argument("--output", required=True, help="Output .docx file path")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    with open(input_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    generate_docx(data, Path(args.output))


if __name__ == "__main__":
    main()
