"""Generate a Word document summarizing InsightRadar updates in simple Chinese."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# ── Title ──
title = doc.add_heading("InsightRadar 更新总结", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

doc.add_paragraph(
    "版本 0.18.0 → 0.20.0  |  2026-02-20  |  共完成 27 项优化"
).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

# ── 一句话总结 ──
doc.add_heading("这次更新做了什么？", level=1)
p = doc.add_paragraph()
p.add_run("简单来说：").bold = True
p.add_run("让系统更安全、更快、更好维护、更容易部署。就像给房子做了一次全面装修——"
           "加固了地基、装了监控、换了更好的锁、还配了一键搬家工具。")

doc.add_paragraph("")

# ── 四大类更新 ──
doc.add_heading("更新分四大类", level=1)

# ── Tier 1 ──
doc.add_heading("1. 安全和基础加固（7 项）", level=2)
items_t1 = [
    ("修好了过时的时间代码", "就像把家里的老式钟换成电波钟，时间更准确，不会出错。"),
    ("加了输入长度限制", "防止有人往搜索框里输入超长内容来搞破坏，相当于给门口装了安检。"),
    ("加了安全防护头", "浏览器访问时多了一层保护，防止恶意代码注入。"),
    ("加了跨域保护", "允许其他网站安全地调用我们的接口。"),
    ("加了网页缓存", "热门数据 60 秒内重复访问不用重新查数据库，网页打开更快。"),
    ("加了 API 文档", "打开 /docs 页面就能看到所有接口的说明书，方便开发者使用。"),
    ("数据库查询加了上限", "防止一次查太多数据把系统拖慢。"),
]
for title_text, desc in items_t1:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(title_text + "：").bold = True
    p.add_run(desc)

doc.add_paragraph("")

# ── Tier 2 ──
doc.add_heading("2. 代码质量提升（7 项）", level=2)
items_t2 = [
    ("代码整理拆分", "原来一个 800 多行的大文件，拆成了 4 个小文件，"
     "就像把一个巨大的衣柜换成了几个分类收纳柜，找东西更方便。"),
    ("统一的错误提示", "出错时返回统一格式的提示信息，不再是乱七八糟的报错。"),
    ("更好的配置管理", "系统设置可以通过环境变量灵活调整，不用改代码。"),
    ("数据库自动管理", "用完自动关闭连接，就像水龙头用完自动关，不浪费资源。"),
    ("日志记录", "系统运行的每一步都有记录，出了问题能快速定位原因。"),
    ("自动化测试", "写了 14 个测试用例，每次改代码都能自动检查有没有改坏。"),
    ("代码风格检查", "提交代码前自动检查格式，保持代码整洁一致。"),
]
for title_text, desc in items_t2:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(title_text + "：").bold = True
    p.add_run(desc)

doc.add_paragraph("")

# ── Tier 3 ──
doc.add_heading("3. 新功能添加（7 项）", level=2)
items_t3 = [
    ("过期数据自动清理", "超过 30 天的旧数据自动删除，数据库不会越来越臃肿。"),
    ("系统健康检查", "访问一个网址就能看到系统是否正常运行，像体检报告一样。"),
    ("实时消息推送", "数据更新时可以实时通知前端页面，不用手动刷新。"),
    ("用户偏好保存", "你的语言、主题等设置会被记住，下次打开还是你喜欢的样子。"),
    ("一键导出数据", "可以把所有数据下载成 JSON 或 CSV 文件，方便在 Excel 里查看。"),
    ("RSS 源健康监控", "能看到每个新闻来源是否正常工作、响应速度如何。"),
    ("全文搜索", "像百度一样，输入关键词就能搜索所有文章内容，又快又准。"),
]
for title_text, desc in items_t3:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(title_text + "：").bold = True
    p.add_run(desc)

doc.add_paragraph("")

# ── Tier 4 ──
doc.add_heading("4. 架构升级（6 项）", level=2)
items_t4 = [
    ("异步数据库", "原来查数据库时系统会'卡住'等结果，"
     "现在可以同时处理多个请求，就像从单车道变成了多车道。"),
    ("内存缓存", "热门数据存在内存里，下次访问直接读内存，比查数据库快几十倍。"),
    ("Docker 容器化", "一个命令就能把整个系统打包运行，"
     "就像把所有行李装进一个箱子，搬到哪里都能马上用。"),
    ("CI/CD 自动化", "每次提交代码，GitHub 会自动检查代码质量和运行测试，"
     "通过了才能合并，防止有问题的代码上线。"),
    ("性能监控", "自动记录每个请求花了多少时间，哪里慢一目了然。"),
    ("访问限流", "限制每分钟的请求次数（比如 AI 搜索最多 10 次/分钟），"
     "防止有人恶意刷接口拖垮系统。"),
]
for title_text, desc in items_t4:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(title_text + "：").bold = True
    p.add_run(desc)

doc.add_paragraph("")

# ── Summary Table ──
doc.add_heading("总览表", level=1)

table = doc.add_table(rows=5, cols=4, style="Light Grid Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["类别", "内容", "数量", "状态"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

rows_data = [
    ["安全加固", "时间修复 / 输入校验 / 安全头 / 缓存 / API 文档", "7 项", "全部完成"],
    ["代码质量", "代码拆分 / 错误处理 / 配置 / 日志 / 测试 / 规范", "7 项", "全部完成"],
    ["新功能", "健康检查 / 实时推送 / 全文搜索 / 导出 / 清理", "7 项", "全部完成"],
    ["架构升级", "异步DB / 缓存 / Docker / CI-CD / 监控 / 限流", "6 项", "全部完成"],
]
for i, row_data in enumerate(rows_data):
    for j, val in enumerate(row_data):
        table.rows[i + 1].cells[j].text = val

doc.add_paragraph("")

# ── Before / After ──
doc.add_heading("升级前 vs 升级后", level=1)

table2 = doc.add_table(rows=8, cols=3, style="Light Grid Accent 1")
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ["方面", "升级前", "升级后"]
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

comparison = [
    ["安全性", "基本没有防护", "安全头 + 输入校验 + 限流"],
    ["代码结构", "800 行挤一个文件", "4 个模块各司其职"],
    ["测试", "没有自动测试", "14 个测试用例自动运行"],
    ["部署", "手动安装依赖运行", "Docker 一键启动"],
    ["性能", "同步阻塞查询", "异步 + 缓存，快几十倍"],
    ["监控", "出问题才知道", "实时日志 + 性能追踪"],
    ["数据管理", "数据只增不减", "30 天自动清理 + 一键导出"],
]
for i, row_data in enumerate(comparison):
    for j, val in enumerate(row_data):
        table2.rows[i + 1].cells[j].text = val

doc.add_paragraph("")

# ── Footer ──
footer = doc.add_paragraph("InsightRadar v0.20.0 — 全球创新与开源情报聚合系统")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer.runs:
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──
output_path = "docs/InsightRadar-更新总结.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
